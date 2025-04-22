import requests
import json
import time
import os
import re
from docx.shared import Inches
import sys
import docx
from docx.document import Document
# NOTE: Importing specific XML types for isinstance checks
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Row, _Cell # _Cell is needed for type hinting set_cell_text
from docx.text.paragraph import Paragraph
from urllib.parse import urlparse, parse_qs, urlencode, urlunparse

# --- Google Gemini Setup ---
try:
    from google import genai
    from google.genai import types
    # IMPORTANT: Use environment variables or a more secure method for API keys
    GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "YOUR_API_KEY_HERE") # Replace fallback with your key if not using env var
    if GEMINI_API_KEY == "YOUR_API_KEY_HERE":
        print("Warning: GEMINI_API_KEY not set as environment variable. Using placeholder.")
    # Configure the client (adjust model and options as needed)
    gemini_model = genai.Client(api_key="AIzaSyA_OdmrUqEp1levy3LhwjAI1m-d7nU6hSE", http_options={'api_version':'v1alpha'})
    print(f"Gemini AI client initialized.") # Model used: {gemini_model.model_name}") # Check attribute if needed

except ImportError:
    print("Error: Google GenAI library not installed. pip install google-generativeai")
    gemini_model = None # Set model to None
except Exception as e:
    print(f"Error initializing Gemini AI client: {e}")
    gemini_model = None # Set model to None

# --- Force UTF-8 ---
try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except AttributeError:
    print("Warning: Could not reconfigure stdout/stderr encoding (not supported on this platform/version?).")


# --- TestRail API Client (from testCaseGetter.py) ---
class TestRailAPIClient:
    def __init__(self, url: str, email: str, api_key: str):
        if not url.strip().startswith(('http://', 'https://')):
            raise ValueError(f"Invalid TestRail URL: '{url}'. Must start with http:// or https://")
        self.base_url = url.strip().rstrip('/')
        self.email = email
        self.api_key = api_key
        self.session = requests.Session()
        self.session.auth = (email, api_key)
        self.session.headers.update({'Content-Type': 'application/json'})
        print(f"TestRailAPIClient initialized for URL: {self.base_url}")

    def _make_request(self, method: str, full_url: str, **kwargs) -> requests.Response | None:
        """Internal helper to make requests. Returns Response object on success, None on error."""
        try:
            timeout = kwargs.pop('timeout', 60)
            response = self.session.request(method, full_url, timeout=timeout, **kwargs)
            response.raise_for_status()
            return response
        except requests.exceptions.HTTPError as http_err:
            status_code = http_err.response.status_code
            try:
                # Try to decode error response for better logging
                response_text = http_err.response.json()
            except json.JSONDecodeError:
                 response_text = http_err.response.text[:200].encode(sys.stdout.encoding, errors='replace').decode(sys.stdout.encoding)

            error_map = {
                400: "Info: Received 400 Bad Request. Check endpoint/parameters.",
                401: "Error: Received 401 Unauthorized. Check email/API key.",
                403: "Error: Received 403 Forbidden. Check permissions or API key.",
                404: "Info: Received 404 Not Found. Resource likely doesn't exist.",
                429: "Error: Received 429 Too Many Requests (Rate Limit). Consider adding delays."
            }
            message = error_map.get(status_code, f"HTTP error during request.")
            print(f"{message} URL: {method} {full_url} - Status: {status_code}. Response: {response_text}...")
            return None
        except requests.exceptions.Timeout:
            print(f"Error: Request timed out for {method} {full_url}")
            return None
        except requests.exceptions.RequestException as e:
            print(f"Error during network request to {method} {full_url}: {e}")
            return None

    def send_get_request(self, endpoint: str):
        """Sends a GET request for single objects (like a specific test case)."""
        print(f"--- Sending SINGLE GET request ---")
        clean_endpoint = endpoint.lstrip('/')
        full_url = f"{self.base_url}/index.php?/api/v2/{clean_endpoint}"
        print(f"  Target: {full_url}")
        response = self._make_request('GET', full_url)

        if response:
            try:
                 if response.status_code == 204 or not response.content:
                      print(f"  Success: Received status {response.status_code} with no content.")
                      return None
                 try:
                     response_data = response.json()
                 except json.JSONDecodeError:
                     print("  Warning: Initial JSON decode failed. Trying UTF-8 decoding.")
                     response_data = json.loads(response.content.decode('utf-8'))

                 print(f"  Success: Received single object data.")
                 return response_data
            except (json.JSONDecodeError, UnicodeDecodeError) as decode_err:
                print(f"  Error: Could not decode response from {response.url}. Status: {response.status_code}. Error: {decode_err}")
                print(f"  Response text (first 500 bytes): {response.content[:500]}...")
                return None
        else:
            print(f"  Failed to get response for single request.")
            return None

# --- Helper Functions ---

def set_cell_text(cell: _Cell, text: str):
    """ Safely sets the text in a cell, replacing existing content. Handles multi-line text. """
    if not isinstance(cell, _Cell):
        print(f"  Error: Invalid object passed to set_cell_text. Expected a Cell, got {type(cell)}.")
        return
    try:
        # Clear existing content more robustly
        cell.text = "" # Clear simple text
        # Remove all but the first paragraph
        while len(cell.paragraphs) > 1:
            p_to_remove = cell.paragraphs[-1]._element
            if p_to_remove.getparent() is not None:
                p_to_remove.getparent().remove(p_to_remove)
            else: # Cannot find parent, maybe already removed or structure issue
                break

        # Ensure at least one paragraph exists
        if not cell.paragraphs:
             cell.add_paragraph()

        # Clear runs in the first paragraph
        first_para = cell.paragraphs[0]
        for run in first_para.runs:
             run.clear()
        first_para.text = "" # Ensure text is cleared too

        # Add the new text, handling newlines as paragraph breaks
        lines = str(text).split('\n')
        first_para.text = lines[0]
        for line in lines[1:]:
            # Add subsequent lines as new paragraphs within the same cell
            cell.add_paragraph(line)

    except Exception as e:
        print(f"  Warn: Exception setting cell text (text: '{str(text)[:50]}...'): {e}")
        # Fallback just in case
        cell.text = str(text) # Direct assignment as last resort

def parse_setup_equipment(setup_text: str) -> list[str]:
    """ Extracts equipment names from the custom_tc_setup string (improved). """
    equipment = []
    if not setup_text:
        return []

    lines = setup_text.splitlines()
    in_equipment_section = False
    equipment_line_regex = re.compile(r"^(?:\|\||>\s*|-\s*|\*\s*|)(?:.*?\|)?\s*([^|#\n]+?)\s*(?:\|.*)?$", re.IGNORECASE)
    excluded_keywords = {'refer to protocol', 'as needed', 'equipment/material', 'quantity', 'part/list number', 'notes:', ''}

    for line in lines:
        line = line.strip()
        if '#**Equipment & Materials**#' in line:
            in_equipment_section = True
            continue
        if in_equipment_section and line.startswith('#**'):
            in_equipment_section = False
            break
        if not in_equipment_section or not line:
            continue

        match = equipment_line_regex.match(line)
        if match:
            potential_name = match.group(1).strip().rstrip('.')
            cleaned_name = ' '.join(potential_name.split())
            if cleaned_name.lower() not in excluded_keywords and len(cleaned_name) > 1:
                 if not ('quantity' in cleaned_name.lower() and 'part' in cleaned_name.lower()):
                    equipment.append(cleaned_name)

    unique_equipment = list(dict.fromkeys(equipment))
    return unique_equipment

# Note: extract_requirement_text_from_expected is not used if AI handles summary text extraction
# Keep it if needed as a fallback or for direct parsing approaches.

def get_gemini_analysis(test_case_data: dict) -> dict | None:
    """ Sends test case data to Gemini for analysis and returns the structured JSON response. """
    if not gemini_model:
        print("Error: Gemini AI model not available.")
        return None
    if not test_case_data:
        print("Error: No test case data provided for Gemini analysis.")
        return None

    type_id_map = { 6: "Functional", 8: "Performance", 7: "Inspection" }
    test_type = type_id_map.get(test_case_data.get('type_id'), "Unknown")

    prompt_data = {
        "test_case_id": test_case_data.get("id"),
        "title": test_case_data.get("title"),
        "requirements": test_case_data.get("refs"),
        "test_type_id": test_case_data.get("type_id"),
        "identified_test_type": test_type,
        "setup_section": test_case_data.get("custom_tc_setup"),
        "steps_expected": test_case_data.get("custom_steps_separated", [])
    }

    # System prompt remains the same as before
    system_prompt = """You are an expert Hardware Verification Engineer specializing in medical devices, specifically ICU Medical infusion pumps and TestRail documentation. Your task is to analyze TestRail test case data and determine the structure needed for a corresponding '.docx' Record File used during test execution.

    Focus on identifying:
    1.  **Test Type:** Confirm the likely test type (Functional, Performance, Inspection) based on the provided 'identified_test_type'.
    2.  **Equipment List:** Extract the names of all necessary equipment from the 'setup_section'.
    3.  **General Information:** Extract the Test Case ID (prefix with 'C'), Test ID (often in refs or custom_test_id), and formulate the main title line combining ID, title, and refs.
    4.  **Results Table Structure (`results_tables` list):** This is critical. Analyze the 'steps_expected' ('content' and 'expected') to determine the necessary results tables and their column headers.
        *   **Identify Recording Actions:** Find *every* instance requiring data recording or verification per sample (e.g., "Record...", "Verify...", "Check...", "Time recorded...").
        *   **Determine Columns:** For each table:
            *   Start with `"No."`.
            *   Add columns for essential identifiers mentioned in setup or steps that need recording per sample (e.g., "Infuser Serial Number", "Battery Serial Number", "Flatbed Serial Number", "Stopwatch Cal_ID / Calibration Due Date"). Determine these dynamically based on the test case context.
            *   Add a distinct column for *each* condition, check, or measurement identified in the recording actions.
            *   End with `"Sign Name/Date"`.
        *   **Header Formatting:**
            *   For standard columns like "No.", device serial numbers, calibration info, and "Sign Name/Date", the header string in the JSON list should be **single-line** (e.g., `"No."`, `"Battery Serial Number"`, `"Sign Name/Date"`).
            *   For condition/check/measurement columns, the header string in the JSON list MUST be **three lines separated by '\\n'**:
                *   Line 1: Concise prose description (e.g., "Time Charging Current Dropped to 0mA (hh:mm:ss)", "Verify Module Presence (Pass/Fail)"). Include units for measurements. Use "(Pass/Fail)" only if the direct result recorded IS pass/fail.
                *   Line 2: Step reference (e.g., "Step 1 Sub-step 8"). Extract from the corresponding step content.
                *   Line 3: Requirement ID and type (e.g., "CADD-SYSDI-73 (condition)", "CADD-HWDI-EC-137 (verification)").
            *   **Example AI Header List:** `["No.", "Infuser Serial Number", "Time Dropped (s)\\nStep 2 Sub-step 5\\nCADD-REQ-101 (condition)", "Verify Alarm (Pass/Fail)\\nStep 2 Sub-step 6\\nCADD-REQ-101 (verification)", "Sign Name/Date"]`
        *   **Table Splitting:** If the total number of columns (including standard/equipment/conditions) exceeds a threshold (e.g., 10 columns), split the *condition* columns into multiple tables.
            *   Generate a separate entry in the `results_tables` list for each split table.
            *   Each split table MUST contain `"No."`, relevant equipment ID column(s), the subset of condition columns, and `"Sign Name/Date"`.
            *   Adjust the table titles accordingly (e.g., "TABLE 1. REQUIREMENTS ... PART 1", "TABLE 2. REQUIREMENTS ... PART 2").
        *   **Table Title:** Generate an appropriate title for each results table (e.g., "TABLE 1. REQUIREMENT [REQ_ID] CONDITION VERIFICATION" or including multiple Req IDs if applicable).
    5.  **Requirements Summary:** Analyze the 'expected' fields. For each requirement ID listed in 'requirements', extract the core verification condition text (the part after "...as a 'Pass' if..."). **Do NOT include step/sub-step numbers in the `verification_text`**. Create a list of objects, each containing `req_id`, `verification_text`, and the associated `step` and `sub_step` numbers (as separate fields).

    **Output Format:** Return the analysis strictly as a JSON object:
    ```json
    {
      "test_type": "Functional | Performance | Inspection | Unknown",
      "general_info": {
        "title_line": "string",
        "test_case_id": "string", // Number only, code will add 'C'
        "test_id": "string"
      },
      "equipment": ["string", ...],
      "results_tables": [
        {
          "title": "string", // Title for this specific table
          "headers": ["string", ...], // List of header strings (single or multi-line as specified)
          "num_rows": number // Functional=29, Performance=30, Inspection=1+
        },
        // Add more table objects here if splitting is needed
        ...
      ],
      "requirements_summary": [
        {
          "step": "string", // "N/A" if not found
          "sub_step": "string", // "N/A" if not found
          "req_id": "string",
          "verification_text": "string" // Core condition text ONLY
        },
        ...
      ]
    }
    ```
    Be precise. Ensure header format rules (single vs. multi-line) are followed. Dynamically determine equipment columns. Split tables correctly if needed.
    """

    user_content = f"""Analyze the following TestRail test case data and generate the JSON structure for its corresponding record file:

    ```json
    {json.dumps(prompt_data, indent=2)}
    ```
    """

    print("--- Sending request to Gemini AI for analysis ---")

    try:
        # Use the generate_content method of the model object
        response = gemini_model.models.generate_content(
            model="gemini-2.5-pro-exp-03-25",
            config=types.GenerateContentConfig(system_instruction=system_prompt),
            contents=user_content
             # system_instruction=system_prompt # Pass system prompt if supported this way by your SDK version
        )

        # Access response text
        if hasattr(response, 'text'):
            response_text = response.text
        elif hasattr(response, 'parts') and response.parts:
             response_text = "".join(part.text for part in response.parts)
        elif hasattr(response, 'candidates') and response.candidates: # Handle candidate structure
              try:
                  response_text = response.candidates[0].content.parts[0].text
              except (IndexError, AttributeError) as e:
                   print(f"Error accessing text through candidates structure: {e}")
                   print(f"Raw Gemini Response: {response}")
                   return None
        else:
             print("Error: Could not extract text from Gemini response. Structure unexpected.")
             print(f"Raw Gemini Response: {response}")
             return None

        # Clean the response text
        cleaned_response_text = re.sub(r'^```json\s*', '', response_text.strip(), flags=re.IGNORECASE)
        cleaned_response_text = re.sub(r'\s*```$', '', cleaned_response_text)

        print("--- Received Gemini AI analysis ---")
        # print(cleaned_response_text) # Debug

        analysis_data = json.loads(cleaned_response_text)
        print("  Successfully parsed Gemini JSON response.")
        return analysis_data

    except json.JSONDecodeError as json_err:
        print(f"Error: Could not decode JSON response from Gemini AI: {json_err}")
        print(f"Raw Gemini Response Text:\n{response_text}")
        return None
    except Exception as e:
        print(f"Error during Gemini AI request or processing: {e}")
        if hasattr(response, 'prompt_feedback'):
             print(f"Prompt Feedback: {response.prompt_feedback}")
        # print(f"Raw Gemini Response Object: {response}") # More detailed debug
        return None

# --- Document Population Logic ---

class RecordFileGenerator:
    """ Populates a Word template based on TestRail data and AI analysis. """

    def __init__(self, template_path: str, testrail_data: dict, ai_analysis: dict):
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        self.doc = docx.Document(template_path)
        self.testrail_data = testrail_data
        self.ai_analysis = ai_analysis
        print(f"RecordFileGenerator initialized with template: {template_path}")

        # --- Define FIXED Element Indices & Coordinates (Updated based on parser dump) ---
        self.TITLE_PARA_INDEX = 0
        self.GENERAL_INFO_TABLE_INDEX = 0 # The actual index in doc.tables

        # Coordinates within General Info Table (Table at index 1)
        # Adjusted to target the start of visually merged cells based on dump/PDF
        self.TC_ID_CELL = (11, 23) # Row 11, Col 18
        self.TEST_ID_CELL = (12, 23) # Row 12, Col 18

        # --- Constants for finding other tables DYNAMICALLY ---
        self.EQUIPMENT_TABLE_SEARCH_HEADER = "Equipment and Material Description"
        # Using more specific search to avoid accidental matches
        self.RESULTS_TABLE_1_SEARCH_TEXT = "TABLE 1. REQUIREMENTS" # Will likely find element 8 -> table 9
        self.REQUIREMENTS_SUMMARY_TABLE_SEARCH_TEXT = "TABLE 2. REQUIREMENTS" # Will likely find element 10 -> table 11

        # --- Constants for Equipment Table Structure ---
        self.EQUIPMENT_START_DATA_ROW = 1
        self.EQUIPMENT_END_DATA_ROW = 10 # Last *template* row for potential matching
        self.EQUIPMENT_NAME_COL = 0
        self.EQUIPMENT_QTY_COL = 1
        self.EQUIPMENT_PART_COL = 2
        self.EQUIPMENT_LOT_COL = 3
        self.EQUIPMENT_CAL_COL = 4

        # --- Constants for Results Table Structure (Specific to Functional Template's Table 1 (Element 9)) ---
        self.RESULTS_TABLE_DATA_START_ROW = 3 # Data starts at row index 3 (Sample '1')

        # --- Constants for Requirements Summary Table Structure ---
        self.REQ_SUMMARY_START_ROW = 1
        self.REQ_SUMMARY_STEP_COL = -1 # No dedicated column in this template
        self.REQ_SUMMARY_SUBSTEP_COL = -1 # No dedicated column in this template
        self.REQ_SUMMARY_REQID_COL = 0
        self.REQ_SUMMARY_VERIF_COL = 1
        self.REQ_SUMMARY_RESULT_COL = 2
        self.REQ_SUMMARY_SIGN_COL = 3

    def _find_table_by_preceding_paragraph_text(self, search_text_regex: str) -> Table | None:
        """ Finds a table immediately following a paragraph matching the regex pattern. """
        try:
            pattern = re.compile(search_text_regex, re.IGNORECASE)
            # Iterate through the raw body elements to maintain order
            body_elements = list(self.doc.element.body)
            for i, element in enumerate(body_elements):
                 if isinstance(element, CT_P):
                     # Map XML element back to Paragraph object to get text
                     para = None
                     for p in self.doc.paragraphs:
                         if p._element == element:
                             para = p
                             break
                     if para and pattern.search(para.text):
                         # Check if the *next* element is a table
                         if i + 1 < len(body_elements) and isinstance(body_elements[i+1], CT_Tbl):
                             # Map table element back to Table object
                             table_element = body_elements[i+1]
                             for t in self.doc.tables:
                                 if t._element == table_element:
                                     print(f"  -> Found table using preceding paragraph regex: '{search_text_regex}'")
                                     return t
            print(f"  -> Warning: Could not find table using preceding paragraph regex: '{search_text_regex}'")
            return None
        except Exception as e:
            print(f"  -> Error finding table by preceding paragraph regex '{search_text_regex}': {e}")
            return None

    def _find_table_by_header_content(self, header_text_keyword: str) -> Table | None:
        """ Finds a table by checking the content of its first row cells. """
        try:
            for table in self.doc.tables:
                if len(table.rows) > 0:
                    header_row = table.rows[0]
                    for cell in header_row.cells:
                        # Handle potential None text or complex cell content
                        cell_text_parts = []
                        for p in cell.paragraphs:
                            cell_text_parts.append(p.text)
                        cell_text = "".join(cell_text_parts).strip()

                        if header_text_keyword.lower() in cell_text.lower():
                            print(f"  -> Found table using header content keyword: '{header_text_keyword}'")
                            return table
            print(f"  -> Warning: Could not find table using header content keyword: '{header_text_keyword}'")
            return None
        except Exception as e:
            print(f"  -> Error finding table by header keyword '{header_text_keyword}': {e}")
            return None

    def populate_title(self):
        """ Populates the main title paragraph (Element 0). """
        print("--- Populating Title ---")
        try:
            title_text = self.ai_analysis.get('general_info', {}).get('title_line', "[Title Missing]")
            if self.TITLE_PARA_INDEX < len(self.doc.paragraphs):
                para0 = self.doc.paragraphs[self.TITLE_PARA_INDEX]
                # Clear existing runs before setting new text
                for run in para0.runs:
                   run.clear()
                para0.text = title_text # Set text directly
                print(f"  Populated title: '{title_text[:100]}...'")
            else:
                print(f"  Error: Paragraph index {self.TITLE_PARA_INDEX} out of range.")
        except Exception as e:
            print(f"  Error populating title: {e}")

    def populate_header(self):
        """ Populates the Test Case ID placeholder in the document header. """
        print("--- Populating Document Header ---")
        try:
            if not self.doc.sections:
                print("  Warning: Document has no sections. Cannot access header.")
                return

            section = self.doc.sections[0]
            header = section.header

            if not header:
                print("  Warning: Section 0 has no header part.")
                return

            tc_id_raw = self.ai_analysis.get('general_info', {}).get('test_case_id', 'CXXXXX')
            tc_id = f"C{tc_id_raw}" if tc_id_raw != 'CXXXXX' and not str(tc_id_raw).startswith('C') else tc_id_raw # Prepend 'C' if needed
            placeholder = "CXXXXX"
            replacement_done = False

            # Find the specific paragraph containing the placeholder
            target_para = None
            for para in header.paragraphs:
                if placeholder in para.text:
                    target_para = para
                    print(f"  Found target paragraph in header: '{target_para.text[:60]}...'")
                    break # Found the paragraph, stop searching

            if target_para:
                # Perform replacement within the runs of the target paragraph only
                inline = target_para.runs
                for i in range(len(inline)):
                    if placeholder in inline[i].text:
                        text = inline[i].text.replace(placeholder, str(tc_id)) # Use modified tc_id
                        inline[i].text = text
                        replacement_done = True
                        print(f"    Replaced placeholder in run {i}. New text: '{inline[i].text}'")
                        # Optional: break here if you only expect one occurrence per paragraph
                        # break
                # Fallback if runs didn't contain it directly (less likely but possible)
                if not replacement_done and placeholder in target_para.text:
                    print("    Placeholder found in paragraph text but not runs, attempting direct text replacement (may affect formatting).")
                    target_para.text = target_para.text.replace(placeholder, str(tc_id)) # Use modified tc_id
                    replacement_done = True

            if not replacement_done:
                print(f"  Warning: Placeholder '{placeholder}' not found or replaced in header paragraphs.")

        except Exception as e:
            print(f"  Error populating header: {e}")
            import traceback
            traceback.print_exc()
    def populate_general_info(self):
        """ Populates fields in the General Information table (Table at index 1). """
        print("--- Populating General Info Table ---")
        try:
            # Access table by its known index
            if self.GENERAL_INFO_TABLE_INDEX < len(self.doc.tables):
                table0 = self.doc.tables[self.GENERAL_INFO_TABLE_INDEX]
                print(f"  Accessed Table {self.GENERAL_INFO_TABLE_INDEX}. Rows: {len(table0.rows)}, Cols: {len(table0.columns)}") # DEBUG: Print dimensions

                # --- DEBUG: Inspect Row 11 ---
                target_row_index = 11
                if target_row_index < len(table0.rows):
                    print(f"  Inspecting Row {target_row_index}:")
                    row11 = table0.rows[target_row_index]
                    print(f"    Number of cells in row {target_row_index}: {len(row11.cells)}")
                    for c_idx, cell in enumerate(row11.cells):
                        # Get text, handling potential errors
                        try:
                            cell_text = cell.text.strip().replace('\n', ' ')
                        except Exception as e:
                            cell_text = f"[Error reading cell: {e}]"
                        print(f"      Cell({target_row_index}, {c_idx}): '{cell_text[:50]}...'") # Limit text length
                else:
                    print(f"  Error: Row index {target_row_index} is out of bounds for table {self.GENERAL_INFO_TABLE_INDEX}.")
                # --- END DEBUG ---

                gen_info = self.ai_analysis.get('general_info', {})
                tc_id = gen_info.get('test_case_id', 'N/A')
                test_id = gen_info.get('test_id', 'N/A') # Assuming Test ID still goes below TC ID
                tc_id_raw = gen_info.get('test_case_id', 'N/A')
                tc_id = f"C{tc_id_raw}" if tc_id_raw != 'N/A' and not str(tc_id_raw).startswith('C') else tc_id_raw # Prepend 'C' if needed
                # Try writing to the cells again (using the coordinates from __init__)
                # We still *expect* (11, 18) and (12, 18) to be the logical targets
                # The debug output above might give clues if these are wrong
                try:
                    # Make sure row 11 exists before accessing
                    if self.TC_ID_CELL[0] < len(table0.rows):
                        set_cell_text(table0.cell(self.TC_ID_CELL[0], self.TC_ID_CELL[1]), tc_id) # Use the modified tc_id
                        print(f"  Populated Test Case ID at {self.TC_ID_CELL}: {tc_id}")
                    else:
                        print(f"  Error: Row index {self.TC_ID_CELL[0]} out of range for Test Case ID.")
                except IndexError:
                    print(f"  IndexError: Could not write to Cell index {self.TC_ID_CELL} for Test Case ID in Table {self.GENERAL_INFO_TABLE_INDEX}.")

                # Test ID (assuming it goes in row 12, same starting column)
                target_test_id_cell = (self.TEST_ID_CELL[0], self.TEST_ID_CELL[1]) # Use coordinates from __init__
                try:
                    # Make sure row 12 exists
                    if target_test_id_cell[0] < len(table0.rows):
                        set_cell_text(table0.cell(target_test_id_cell[0], target_test_id_cell[1]), test_id)
                        print(f"  Attempted to populate Test ID at {target_test_id_cell}: {test_id}")
                    else:
                        print(f"  Error: Row index {target_test_id_cell[0]} out of range for Test ID.")
                except IndexError:
                    print(f"  IndexError: Could not write to Cell index {target_test_id_cell} for Test ID in Table {self.GENERAL_INFO_TABLE_INDEX}.")

            else:
                print(f"  Error: Table index {self.GENERAL_INFO_TABLE_INDEX} out of range for General Info.")
        except Exception as e:
            print(f"  Error populating General Info table: {e}")
            import traceback
            traceback.print_exc()

    def populate_equipment(self):
        """ Clears existing equipment data rows (if any), populates with AI-identified equipment, and applies a standard table style. """
        print("--- Populating Equipment Table (Replacing Template Rows & Applying Style) ---")
        try:
            # Find the table using header content
            table1 = self._find_table_by_header_content(self.EQUIPMENT_TABLE_SEARCH_HEADER)
            if not table1:
                print("  Error: Equipment table not found using header search.")
                return

            equipment_list = self.ai_analysis.get('equipment', [])
            if not equipment_list:
                # Fallback (optional, keep if desired)
                # print("  Warning: Equipment list not found in AI analysis, attempting to parse from setup text.")
                # equipment_list = parse_setup_equipment(self.testrail_data.get('custom_tc_setup', ''))
                print("  No equipment identified by AI or setup text.")
                # Decide if you still want to clear template rows even if no equipment is found
                # For now, let's just return if no equipment is to be added.
                return # Exit if no equipment

            print(f"  Equipment to populate: {equipment_list}")

            # --- Clear any existing data rows (just in case template wasn't fully cleaned) ---
            header_row_count = 1 # Assuming 1 header row
            while len(table1.rows) > header_row_count:
                row_to_remove = table1.rows[-1]
                row_element = row_to_remove._element
                if row_element.getparent() is not None:
                    row_element.getparent().remove(row_element)
                else:
                    break
            print(f"  Cleared any pre-existing data rows. Table rows now: {len(table1.rows)}")

            # --- Add new rows for each piece of equipment ---
            for eq_name in equipment_list:
                try:
                    new_row = table1.add_row()
                    # Populate only the first column (equipment name)
                    set_cell_text(new_row.cells[self.EQUIPMENT_NAME_COL], eq_name)
                    # Clear other columns in the new row to ensure they are blank
                    for c_idx in range(1, len(new_row.cells)):
                        # Check cell index is valid before accessing
                        if c_idx < len(new_row.cells):
                            set_cell_text(new_row.cells[c_idx], "") # Set to empty string
                        else:
                            print(f"  Warning: Column index {c_idx} out of range when clearing added row.")
                    print(f"  Added row for equipment: '{eq_name}'")
                except IndexError:
                    print(f"  Error: Column index {self.EQUIPMENT_NAME_COL} out of range when adding row for '{eq_name}'.")
                except Exception as add_e:
                    print(f"  Error adding row for equipment '{eq_name}': {add_e}")

            # --- Apply a standard table style AFTER adding all rows ---
            # 'Table Grid' is a basic style that usually includes all borders.
            # Other options: 'Light Shading Accent 1', 'Medium Grid 1 Accent 1', etc.
            # Check Word's Table Design tab for names. Style must exist in the template or Word's defaults.
            try:
                target_style = 'RecordHeaderStyle'
                print(f"  Attempting to apply style '{target_style}' to the equipment table.")
                table1.style = target_style
                print(f"  Applied style '{table1.style.name}' to equipment table.") # Verify applied style name
            except Exception as style_e:
                print(f"  Warning: Could not apply table style '{target_style}'. Error: {style_e}")
                print("  Borders might still be missing. Consider checking template's default table style.")

            print(f"  Finished populating equipment table with {len(equipment_list)} items.")

        except Exception as e:
            print(f"  Error populating Equipment table: {e}")
            import traceback
            traceback.print_exc()

    def populate_results_tables(self):
        """ Populates results tables dynamically, adding columns, setting width/style/formatting. """
        print("--- Populating Results Tables (Dynamic Columns) ---")
        results_tables_data = self.ai_analysis.get('results_tables', [])
        if not results_tables_data:
            print("  No results table structures found in AI analysis.")
            return

        table_placeholders = [
            "PLACEHOLDER_RESULTS_TABLE_1",
            "PLACEHOLDER_RESULTS_TABLE_2",
        ]
        default_column_width = Inches(1.0) # Default width used when adding columns
        target_total_width_inches = 10.0

        for i, table_data in enumerate(results_tables_data):
            if i >= len(table_placeholders):
                print(f"  Warning: AI provided data for results table {i+1}, but only {len(table_placeholders)} placeholders defined/found. Skipping.")
                break

            placeholder_text = table_placeholders[i]
            print(f"\n  Processing Results Table {i+1} (Looking for placeholder '{placeholder_text}')...")

            target_table = None
            title_para = None
            body_elements = list(self.doc.element.body)
            for elem_idx, element in enumerate(body_elements):
                if isinstance(element, CT_P):
                    para = next((p for p in self.doc.paragraphs if p._element == element), None)
                    if para and placeholder_text in para.text:
                        title_para = para
                        if elem_idx + 1 < len(body_elements) and isinstance(body_elements[elem_idx+1], CT_Tbl):
                            table_element = body_elements[elem_idx+1]
                            target_table = next((t for t in self.doc.tables if t._element == table_element), None)
                        break

            if not target_table:
                print(f"  Error: Could not find table following placeholder '{placeholder_text}'. Skipping Table {i+1}.")
                continue

            table_title = table_data.get("title", f"Results Table {i+1}")
            headers = table_data.get("headers", [])
            num_data_rows = table_data.get("num_rows", 1)

            if not headers:
                print(f"    Warning: No headers found for table '{table_title}'. Skipping.")
                continue

            print(f"    Found table. Populating title: '{table_title}' and processing {len(headers)} headers.")

            # --- Set Table Title ---
            if title_para:
                print(f"    Setting title in placeholder paragraph...")
                for run in title_para.runs: run.clear()
                title_para.text = table_title
                if title_para.runs: title_para.runs[0].font.bold = True
            else:
                print(f"    Warning: Placeholder paragraph not found. Cannot set title.")

             # --- Add Columns and Set Widths using EMUs ---
            num_needed_cols = len(headers)
            current_cols = len(target_table.columns)
            num_cols_to_add = num_needed_cols - current_cols

            if num_cols_to_add < 0:
                print(f"    Warning: Template table {i+1} has more columns ({current_cols}) than AI headers ({num_needed_cols}).")
                num_cols_to_add = 0

            # --- Set table layout to fixed BEFORE adjusting widths ---
            try:
                print("    Setting table layout to fixed.")
                target_table.autofit = True
                # Setting overall table width in EMU (optional, but may help)
                # total_table_width_emu = 9144000 # 10 inches in EMU
                # target_table.width = total_table_width_emu
            except Exception as layout_e:
                print(f"    Warning: Could not set table layout properties: {layout_e}")
            # ---

            # --- Calculate target width PER COLUMN in EMUs ---
            calculated_width_per_column_emu = None
            target_total_width_emu = 9144000 # 10 inches * 914400 EMU/inch
            if num_needed_cols > 0:
                try:
                    # Calculate width as an integer EMU value
                    calculated_width_per_column_emu = int(target_total_width_emu / num_needed_cols)
                    print(f"    Calculated width per column: {calculated_width_per_column_emu} EMU (approx {(calculated_width_per_column_emu / 914400):.2f} inches)")
                except ZeroDivisionError:
                    print("    Warning: Cannot divide by zero columns for width calculation.")
            else:
                print("   Warning: Zero columns needed, cannot calculate width.")
            # --- End EMU Calculation ---

            # Default width for adding columns initially (can be small, will be overridden)
            initial_add_width_emu = 100000 # A small default EMU value

            if num_cols_to_add > 0:
                print(f"    Adding {num_cols_to_add} columns and setting widths (EMU)...")
                for col_idx_to_add in range(num_cols_to_add):
                    try:
                        # Add the column using the minimal EMU width
                        new_col_ref = target_table.add_column(initial_add_width_emu)

                        # --- IMMEDIATELY set the calculated EMU width ---
                        current_col_count = len(target_table.columns)
                        new_col_index = current_col_count - 1
                        if calculated_width_per_column_emu is not None:
                            try:
                                target_table.columns[new_col_index].width = calculated_width_per_column_emu
                                # Also set cell width in first row
                                if len(target_table.rows) > 0:
                                    header_row = target_table.rows[0]
                                    if new_col_index < len(header_row.cells):
                                            header_row.cells[new_col_index].width = calculated_width_per_column_emu
                            except IndexError:
                                print(f"      Warning: Index out of bounds setting EMU width for added column {new_col_index}")
                            except Exception as set_w_e:
                                print(f"      Warning: Error setting EMU width for added column {new_col_index}: {set_w_e}")
                        # ---
                    except Exception as add_col_e:
                        print(f"      Error adding column {col_idx_to_add + 1}: {add_col_e}. Stopping.")
                        return

            # --- Set EMU width for the FIRST column ---
            if calculated_width_per_column_emu is not None and len(target_table.columns) > 0:
                try:
                    print(f"    Setting EMU width for initial column 0.")
                    target_table.columns[0].width = calculated_width_per_column_emu
                    if len(target_table.rows) > 0:
                        target_table.rows[0].cells[0].width = calculated_width_per_column_emu
                except IndexError:
                    print("     Warning: Could not set EMU width for initial column 0.")
                except Exception as set_w0_e:
                    print(f"     Warning: Error setting EMU width for initial column 0: {set_w0_e}")
            # ---

            # Verify final column count
            final_cols = len(target_table.columns)
            print(f"    Table now has {final_cols} columns.")
            if final_cols < num_needed_cols:
                print(f"    Error: Failed to add/set sufficient columns. Skipping.")
                continue

            # --- Populate Headers ---
            if len(target_table.rows) < 3:
                print("    Error: Results table template needs >= 3 header rows. Skipping.")
                continue

            header_row_prose = target_table.rows[0]
            header_row_step = target_table.rows[1]
            header_row_req = target_table.rows[2]
            print(f"    Populating headers across {num_needed_cols} columns...")
            for c_idx in range(num_needed_cols):
                if c_idx >= len(header_row_prose.cells): continue # Safety check

                ai_header_text = headers[c_idx]
                is_standard_col = ai_header_text.count('\n') == 0
                try:
                    if is_standard_col:
                        set_cell_text(header_row_prose.cells[c_idx], ai_header_text)
                        if c_idx < len(header_row_step.cells): set_cell_text(header_row_step.cells[c_idx], "")
                        if c_idx < len(header_row_req.cells): set_cell_text(header_row_req.cells[c_idx], "")
                    else:
                        header_parts = ai_header_text.split('\n')
                        prose = header_parts[0] if len(header_parts) > 0 else ""
                        step_ref = header_parts[1] if len(header_parts) > 1 else ""
                        req_ref = header_parts[2] if len(header_parts) > 2 else ""
                        if c_idx < len(header_row_prose.cells): set_cell_text(header_row_prose.cells[c_idx], prose)
                        if c_idx < len(header_row_step.cells): set_cell_text(header_row_step.cells[c_idx], step_ref)
                        if c_idx < len(header_row_req.cells): set_cell_text(header_row_req.cells[c_idx], req_ref)
                except IndexError:
                    print(f"    Error: Cell index {c_idx} out of range populating headers.")

            # --- Apply Header Formatting (Bold) ---
            print("    Applying bold formatting to header rows...")
            num_header_rows = 3
            for r_idx in range(min(num_header_rows, len(target_table.rows))):
                row = target_table.rows[r_idx]
                for c_idx in range(num_needed_cols):
                    if c_idx < len(row.cells):
                        cell = row.cells[c_idx]
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    else: break

            # --- Manage Data Rows ---
            num_template_header_rows = 3
            data_start_row_index = num_template_header_rows
            existing_data_rows = max(0, len(target_table.rows) - num_template_header_rows)
            rows_to_add = num_data_rows - existing_data_rows
            print(f"    Target data rows: {num_data_rows}. Template has: {existing_data_rows}. Need to add: {rows_to_add}")

            if rows_to_add > 0:
                print(f"    Adding {rows_to_add} data rows...")
                for _ in range(rows_to_add): target_table.add_row()
            elif rows_to_add < 0:
                rows_to_remove_count = abs(rows_to_add)
                print(f"    Removing {rows_to_remove_count} extra template data rows...")
                for _ in range(rows_to_remove_count):
                    if len(target_table.rows) > num_template_header_rows:
                        row_element_to_remove = target_table.rows[-1]._element
                        if row_element_to_remove.getparent() is not None:
                            row_element_to_remove.getparent().remove(row_element_to_remove)
                        else: break
                    else: break

            # --- Populate data rows ---
            print(f"    Populating {num_data_rows} data rows...")
            num_cols_to_populate = num_needed_cols
            for r_idx_offset in range(num_data_rows):
                actual_row_index = data_start_row_index + r_idx_offset
                if actual_row_index >= len(target_table.rows): break
                row = target_table.rows[actual_row_index]
                current_row_cols = len(row.cells)
                
                for c_idx in range(min(current_row_cols, num_cols_to_populate)): # Populate up to available cols
                    header_first_line = headers[c_idx].split('\n')[0].lower()
                    cell = row.cells[c_idx]
                    if c_idx == 0 and "no." in header_first_line:
                        set_cell_text(cell, str(r_idx_offset + 1))
                    elif "pass" in headers[c_idx].lower() and "fail" in headers[c_idx].lower():
                        set_cell_text(cell, "Pass ☐ / Fail ☐")
                    elif c_idx == num_needed_cols - 1 and ("sign" in header_first_line or "/date" in header_first_line):
                        set_cell_text(cell, "")
                    else:
                        set_cell_text(cell, "") # Clear other data cells

            # --- Apply Table Style for Borders ---
            try:
                # Use 'Table Grid' for basic borders OR your custom style name
                # target_style = 'RecordHeaderStyle' # If you created a custom style in Word
                target_style = 'RecordHeaderStyle'
                print(f"    Attempting to apply style '{target_style}' to results table {i+1}.")
                target_table.style = target_style
                print(f"    Applied style '{target_table.style.name}'.")
            except Exception as style_e:
                print(f"    Warning: Could not apply table style '{target_style}'. Error: {style_e}")

    def populate_requirements_summary(self):
        """ Populates the Requirements Summary table (Found dynamically) and its preceding title. """
        print("--- Populating Requirements Summary Table ---")
        summary_data = self.ai_analysis.get('requirements_summary', [])

        # --- Find the table FIRST using the preceding paragraph text ---
        target_table = self._find_table_by_preceding_paragraph_text(self.REQUIREMENTS_SUMMARY_TABLE_SEARCH_TEXT)
        if not target_table:
            print(f"  Error: Could not find Requirements Summary table using search text '{self.REQUIREMENTS_SUMMARY_TABLE_SEARCH_TEXT}'.")
            return # Cannot proceed without the table

        # --- Set Table Title (in the paragraph found *just before* the table) ---
        # This logic assumes _find_table_by_preceding_paragraph_text worked correctly
        table_element = target_table._element
        parent = table_element.getparent()
        title_para = None # Initialize title_para
        if parent is not None:
            try:
                table_xml_index = parent.index(table_element)
                if table_xml_index > 0:
                    preceding_element = parent[table_xml_index - 1]
                    if isinstance(preceding_element, CT_P):
                        # Map the preceding XML paragraph element back to a python-docx Paragraph object
                        for p in self.doc.paragraphs:
                            if p._element == preceding_element:
                                title_para = p
                                break
            except ValueError:
                print("    Warning: Could not determine summary table index within parent.") # Handle case where table might not be direct child
            except Exception as find_e:
                print(f"    Warning: Error finding preceding paragraph for summary table title: {find_e}")

        if title_para:
            # Construct the title dynamically using the requirements from AI analysis
            req_ids = [req.get('req_id', 'XX') for req in summary_data] # Get all req IDs
            req_ids_str = " AND ".join(req_ids) if req_ids else "XX" # Join them, handle empty list
            table_title = f"TABLE 2. REQUIREMENTS {req_ids_str} VERIFICATION" # Construct title
            print(f"    Setting title in preceding paragraph: '{table_title}'")
            for run in title_para.runs: run.clear() # Clear existing formatting/text runs
            title_para.text = table_title # Set the new title text
            if title_para.runs: title_para.runs[0].font.bold = True # Optionally make it bold
        else:
            # This warning indicates the title paragraph wasn't found via the expected structure
            print(f"    Warning: No paragraph found directly preceding the summary table using search '{self.REQUIREMENTS_SUMMARY_TABLE_SEARCH_TEXT}'. Cannot set title paragraph automatically.")
        # --- End Title Setting ---

        # Proceed only if there's data to populate
        if not summary_data:
            print("  No requirements summary data found in AI analysis. Skipping row population.")
            # Still attempt cleanup of template rows if necessary
            template_rows_available = len(target_table.rows) - self.REQ_SUMMARY_START_ROW
            if template_rows_available > 0:
                print(f"  Clearing {template_rows_available} template rows in empty summary table.")
                start_clear_row = self.REQ_SUMMARY_START_ROW
                for r_idx in range(start_clear_row, len(target_table.rows)):
                    row = target_table.rows[r_idx]
                    for cell in row.cells:
                        set_cell_text(cell, "")
            return # Exit after handling empty summary_data

        # --- Populate Rows ---
        num_reqs_to_populate = len(summary_data)
        template_rows_available = len(target_table.rows) - self.REQ_SUMMARY_START_ROW
        print(f"  Populating {num_reqs_to_populate} requirements. Template has {template_rows_available} data rows available.")

        for i, req_info in enumerate(summary_data):
            row_index = self.REQ_SUMMARY_START_ROW + i
            row = None # Initialize row

            if row_index < len(target_table.rows):
                row = target_table.rows[row_index]
            else:
                # Add row if needed
                print(f"    Warning: Not enough rows in summary table template for requirement {i+1}. Adding row.")
                try:
                    row = target_table.add_row()
                except Exception as add_e:
                    print(f"    Error adding row to summary table: {add_e}")
                    break # Stop processing summary if adding rows fails

            if row: # Proceed only if row exists or was added successfully
                print(f"    Populating row {row_index} for Req: {req_info.get('req_id', 'N/A')}")
                try:
                    # Get data from AI
                    step_text = req_info.get('step', '') if req_info.get('step', 'N/A') != 'N/A' else ''
                    sub_step_text = req_info.get('sub_step', '') if req_info.get('sub_step', 'N/A') != 'N/A' else ''
                    req_id_text = req_info.get('req_id', '')
                    verif_text = req_info.get('verification_text', '') # Use original verification text

                    # Populate cells using column indices, checking if index is valid (>= 0)
                    # Step/Substep columns are not populated as indices are -1
                    if self.REQ_SUMMARY_REQID_COL >= 0: set_cell_text(row.cells[self.REQ_SUMMARY_REQID_COL], req_id_text)
                    if self.REQ_SUMMARY_VERIF_COL >= 0: set_cell_text(row.cells[self.REQ_SUMMARY_VERIF_COL], verif_text)
                    if self.REQ_SUMMARY_RESULT_COL >= 0: set_cell_text(row.cells[self.REQ_SUMMARY_RESULT_COL], "Pass ☐ / Fail ☐")
                    if self.REQ_SUMMARY_SIGN_COL >= 0: set_cell_text(row.cells[self.REQ_SUMMARY_SIGN_COL], "")

                except IndexError:
                    print(f"    Error: Column index out of range while populating summary row {row_index}.")
                except Exception as cell_e:
                    print(f"    Error populating cells in summary row {row_index}: {cell_e}")


        # --- Clear Extra Template Rows ---
        if num_reqs_to_populate < template_rows_available:
            rows_to_clear_count = template_rows_available - num_reqs_to_populate
            print(f"  Clearing {rows_to_clear_count} extra rows in summary table template.")
            start_clear_row = self.REQ_SUMMARY_START_ROW + num_reqs_to_populate
            # Iterate backwards when removing to avoid index issues
            for r_idx in range(len(target_table.rows) - 1, start_clear_row - 1, -1):
                row = target_table.rows[r_idx]
                for cell in row.cells:
                    set_cell_text(cell, "") # Clear content of extra rows

        # --- Final Cleanup: Remove any rows beyond what was populated ---
        # This loop ensures the table has exactly the right number of rows at the end
        final_expected_rows = self.REQ_SUMMARY_START_ROW + num_reqs_to_populate
        while len(target_table.rows) > final_expected_rows:
            print(f"  Removing extra trailing row {len(target_table.rows)} from summary table.")
            row_element_to_remove = target_table.rows[-1]._element
            if row_element_to_remove.getparent() is not None:
                row_element_to_remove.getparent().remove(row_element_to_remove)
            else:
                print("    Warning: Could not remove extra summary row, parent not found.")
                break # Avoid infinite loop

    def save(self, output_path: str):
        """ Saves the populated document. """
        print(f"--- Saving populated document to: {output_path} ---")
        try:
            output_dir = os.path.dirname(output_path)
            if output_dir: os.makedirs(output_dir, exist_ok=True)
            self.doc.save(output_path)
            print("  Save successful.")
        except PermissionError:
             print(f"  Error: Permission denied saving to '{output_path}'. Check if the file is open or permissions are correct.")
        except Exception as e:
            print(f"  Error saving document: {e}")
            import traceback
            traceback.print_exc()

# --- Main Execution ---
def main():
    # --- Configuration ---
    TESTRAIL_URL = "https://icumed.testrail.io"
    # IMPORTANT: Replace with your actual credentials (or use environment variables)
    TESTRAIL_EMAIL = os.environ.get("TESTRAIL_EMAIL", "jairo.rodriguez@icumed.com")
    TESTRAIL_API_KEY = os.environ.get("TESTRAIL_API_KEY", "uttDDNwG5EMQeiW71KaN-3l9Ndf4mBqWT0xV7TF5F")

    # test_case_id_to_process = 1043745 # Functional Example 2
    test_case_id_to_process = 1043779 # Functional Example 2

    # Template file paths (Adjust as necessary)
    template_dir = r'C:\Users\JR00093781\OneDrive - ICU Medical Inc\Documents\TestRecordGenerator\v2\Templates'
    template_file_functional = os.path.join(template_dir, 'Functional_Template_Base.docx')
    template_file_performance = os.path.join(template_dir, 'Performance_Template_Base.docx') # Assuming it exists
    template_file_inspection = os.path.join(template_dir, 'Inspection_Template_Base.docx') # Assuming it exists

    # Select template based on test type (Manual selection for now)
    # TODO: Add logic to select template based on ai_analysis['test_type'] later
    template_file_to_use = template_file_functional

    output_dir = r"C:\Users\JR00093781\OneDrive - ICU Medical Inc\Documents\TestRecordGenerator\v2\Generated Record Files"
    safe_tc_id = str(test_case_id_to_process)
    output_file_name = os.path.join(output_dir, f"Record File TC {safe_tc_id} - Generated.docx")

    # --- Initial Checks ---
    if not gemini_model:
        print("Error: Gemini AI Client failed to initialize. Exiting.")
        sys.exit(1)
    if "YOUR_TESTRAIL_KEY" in TESTRAIL_API_KEY or "your_email@example.com" in TESTRAIL_EMAIL:
         print("Error: Please configure TestRail credentials.")
         # sys.exit(1) # Comment out exit for testing if using hardcoded values
    if not os.path.exists(template_file_to_use):
         print(f"Error: Selected template file not found: {template_file_to_use}")
         sys.exit(1)

    # --- Processing ---
    print(f"--- Starting Record File Generation for Test Case ID: {test_case_id_to_process} ---")

    tr_client = TestRailAPIClient(TESTRAIL_URL, TESTRAIL_EMAIL, TESTRAIL_API_KEY)
    endpoint = f"get_case/{test_case_id_to_process}"
    testrail_data = tr_client.send_get_request(endpoint)

    if not testrail_data:
        print(f"Error: Failed to fetch TestRail data for TC {test_case_id_to_process}. Exiting.")
        sys.exit(1)

    ai_analysis = get_gemini_analysis(testrail_data)

    if not ai_analysis:
        print("Error: Failed to get analysis from Gemini AI. Exiting.")
        sys.exit(1)

    print("\n--- Received AI Analysis (Structure) ---")
    print(json.dumps(ai_analysis, indent=2))

    try:
        generator = RecordFileGenerator(template_file_to_use, testrail_data, ai_analysis)
        generator.populate_header()
        generator.populate_title()
        generator.populate_general_info()
        generator.populate_equipment()
        generator.populate_results_tables()
        generator.populate_requirements_summary()
        generator.save(output_file_name)
        print(f"\n--- Record File Generation Completed for TC {test_case_id_to_process} ---")
        print(f"Output saved to: {output_file_name}")
    except FileNotFoundError as fnf_err:
         print(f"\nError initializing generator: {fnf_err}")
    except Exception as gen_err:
         print(f"\nAn critical error occurred during record file generation: {gen_err}")
         import traceback
         traceback.print_exc()

if __name__ == "__main__":
    main()