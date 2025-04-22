import docx
import sys
import os
import re # For finding table

# Force UTF-8
try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except AttributeError: pass

def find_table_by_preceding_paragraph_text(doc, search_text_regex: str):
    # Simplified version for this script
    pattern = re.compile(search_text_regex, re.IGNORECASE)
    body_elements = list(doc.element.body)
    for i, element in enumerate(body_elements):
         if isinstance(element, docx.oxml.text.paragraph.CT_P):
             para = None
             for p in doc.paragraphs:
                 if p._element == element: para = p; break
             if para and pattern.search(para.text):
                 if i + 1 < len(body_elements) and isinstance(body_elements[i+1], docx.oxml.table.CT_Tbl):
                     table_element = body_elements[i+1]
                     for t in doc.tables:
                         if t._element == table_element: return t
    return None

def inspect_table_headers(docx_path, table_search_text):
    try:
        doc = docx.Document(docx_path)
        print(f"--- Inspecting Headers of Table Preceded by '{table_search_text}' ---")

        target_table = find_table_by_preceding_paragraph_text(doc, table_search_text)

        if not target_table:
            print("Error: Target table not found.")
            return

        num_header_rows_to_inspect = 3 # Inspect first 3 rows
        if len(target_table.rows) < num_header_rows_to_inspect:
            print(f"Warning: Table has fewer than {num_header_rows_to_inspect} rows.")
            num_header_rows_to_inspect = len(target_table.rows)

        for r_idx in range(num_header_rows_to_inspect):
            row = target_table.rows[r_idx]
            print(f"\n-- Row {r_idx} --")
            print(f"  Number of logical cells reported by python-docx: {len(row.cells)}")
            for c_idx, cell in enumerate(row.cells):
                try:
                    cell_text = cell.text.strip().replace('\n', ' \\n ') # Show newlines
                except Exception as e:
                    cell_text = f"[Error reading cell: {e}]"
                print(f"    Cell({r_idx}, {c_idx}): '{cell_text[:100]}...'") # Limit output length

    except FileNotFoundError:
        print(f"Error: File not found at {docx_path}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

# --- Usage ---
if __name__ == "__main__":
    # <<< IMPORTANT: Use the path to your TEMPLATE file >>>
    template_path = r'C:\Users\JR00093781\OneDrive - ICU Medical Inc\Documents\TestRecordGenerator\v2\Generated Record Files\Record File TC 1013408 - Generated.docx'
    # The text used to find the results table
    search_text = "TABLE 1. REQUIREMENT"

    if os.path.exists(template_path):
        inspect_table_headers(template_path, search_text)
    else:
        print(f"Error: Template file '{template_path}' does not exist.")